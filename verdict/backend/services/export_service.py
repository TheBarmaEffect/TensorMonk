"""Export service for generating verdict reports in multiple formats."""

import io
import json
from datetime import datetime


def generate_markdown_report(session_data: dict) -> str:
    """Generate a comprehensive markdown report from session data."""
    decision = session_data.get("decision", {})
    research = session_data.get("research_package", {})
    prosecutor = session_data.get("prosecutor_argument", {})
    defense = session_data.get("defense_argument", {})
    witnesses = session_data.get("witness_reports", [])
    verdict = session_data.get("verdict", {})
    synthesis = session_data.get("synthesis", {})
    output_format = session_data.get("output_format", "executive")
    domain = session_data.get("domain", "business")

    lines = []
        if not session_data:
        return "# VERDICT — No Data Available\n\nSession has no results yet."
    lines.append("# VERDICT — AI Courtroom Analysis Report")
    lines.append(f"\n**Decision:** {decision.get('question', 'N/A')}")
    lines.append(f"**Date:** {datetime.utcnow().strftime('%B %d, %Y')}")
    lines.append(f"**Session ID:** {decision.get('id', 'N/A')}")
    lines.append(f"**Output Format:** {output_format.title()}")
    lines.append(f"**Domain:** {domain.title()}")
    lines.append("\n---\n")

    lines.append("## 1. Research Analysis")
    if research:
        if research.get("summary"):
            lines.append(f"\n{research['summary']}\n")
        if research.get("key_data_points"):
            lines.append("### Key Data Points")
            for point in research["key_data_points"]:
                lines.append(f"- {point}")
        if research.get("risk_landscape"):
            lines.append("\n### Risk Landscape")
            for risk in research["risk_landscape"]:
                lines.append(f"- {risk}")

    lines.append("\n---\n")

    lines.append("## 2. Prosecution Arguments (FOR)")
    if prosecutor:
        if prosecutor.get("opening"):
            lines.append(f"\n> {prosecutor['opening']}\n")
        if prosecutor.get("claims"):
            for i, claim in enumerate(prosecutor["claims"], 1):
                lines.append(f"### Claim {i}")
                lines.append(f"**Statement:** {claim.get('statement', '')}")
                lines.append(f"**Evidence:** {claim.get('evidence', '')}")
                lines.append(f"**Confidence:** {round(claim.get('confidence', 0.5) * 100)}%\n")
        lines.append(f"**Overall Confidence:** {round(prosecutor.get('confidence', 0.5) * 100)}%")

    lines.append("\n---\n")

    lines.append("## 3. Defense Arguments (AGAINST)")
    if defense:
        if defense.get("opening"):
            lines.append(f"\n> {defense['opening']}\n")
        if defense.get("claims"):
            for i, claim in enumerate(defense["claims"], 1):
                lines.append(f"### Claim {i}")
                lines.append(f"**Statement:** {claim.get('statement', '')}")
                lines.append(f"**Evidence:** {claim.get('evidence', '')}")
                lines.append(f"**Confidence:** {round(claim.get('confidence', 0.5) * 100)}%\n")
        lines.append(f"**Overall Confidence:** {round(defense.get('confidence', 0.5) * 100)}%")

    lines.append("\n---\n")

    if witnesses:
        lines.append("## 4. Witness Reports")
        for i, w in enumerate(witnesses, 1):
            verdict_emoji = "✅" if w.get("verdict_on_claim") == "sustained" else "❌" if w.get("verdict_on_claim") == "overruled" else "⚠️"
            lines.append(f"\n### Witness {i} ({w.get('witness_type', 'fact').title()}) {verdict_emoji}")
            lines.append(f"**Verdict:** {w.get('verdict_on_claim', 'inconclusive').upper()}")
            lines.append(f"**Resolution:** {w.get('resolution', '')}")
            lines.append(f"**Confidence:** {round(w.get('confidence', 0.5) * 100)}%")
        lines.append("\n---\n")

    lines.append("## 5. Final Verdict")
    if verdict:
        ruling = verdict.get("ruling", "conditional").upper()
        lines.append(f"\n### Ruling: {ruling}")
        lines.append(f"**Confidence:** {round(verdict.get('confidence', 0.5) * 100)}%\n")
        lines.append(f"{verdict.get('reasoning', '')}\n")
        if verdict.get("key_factors"):
            lines.append("### Key Factors")
            for i, factor in enumerate(verdict["key_factors"], 1):
                lines.append(f"{i}. {factor}")

    lines.append("\n---\n")

    lines.append("## 6. Battle-Tested Synthesis")
    if synthesis:
        lines.append(f"\n**Strength Score:** {round(synthesis.get('strength_score', 0.7) * 100)}%\n")
        if synthesis.get("improved_idea"):
            lines.append(synthesis["improved_idea"])
        if synthesis.get("addressed_objections"):
            lines.append("\n### Objections Addressed")
            for obj in synthesis["addressed_objections"]:
                lines.append(f"- ✓ {obj}")
        if synthesis.get("recommended_actions"):
            lines.append("\n### Recommended Actions")
            for i, action in enumerate(synthesis["recommended_actions"], 1):
                lines.append(f"{i}. {action}")

    # Section 7: Verdict Stability Analysis (if available)
    analysis = session_data.get("analysis", {})
    stability = analysis.get("verdict_stability") or session_data.get("verdict_stability")
    quality = analysis.get("argument_quality") or session_data.get("argument_quality")

    if stability or quality:
        lines.append("\n---\n")
        lines.append("## 7. Analysis & Quality Metrics")

        if stability:
            robustness = stability.get("combined_robustness", 0)
            is_robust = stability.get("verdict_is_robust", True)
            margin = stability.get("evidence_margin", "unknown")
            lines.append(f"\n### Verdict Stability")
            lines.append(f"- **Robustness Score:** {round(robustness * 100)}%")
            lines.append(f"- **Verdict Robust:** {'Yes' if is_robust else '⚠️ No — verdict may flip under evidence perturbation'}")
            lines.append(f"- **Evidence Margin:** {margin}")
            flip_rate = stability.get("flip_rate", 0)
            if flip_rate > 0:
                lines.append(f"- **Flip Rate:** {round(flip_rate * 100)}% of perturbation simulations flipped the verdict")

        if quality:
            lines.append(f"\n### Argument Quality")
            pro_q = quality.get("prosecutor", {})
            def_q = quality.get("defense", {})
            if pro_q:
                lines.append(f"- **Prosecution Grade:** {pro_q.get('grade', 'N/A')} ({round(pro_q.get('overall', 0) * 100)}%)")
            if def_q:
                lines.append(f"- **Defense Grade:** {def_q.get('grade', 'N/A')} ({round(def_q.get('overall', 0) * 100)}%)")
            gap = quality.get("quality_gap", 0)
            if gap > 0:
                lines.append(f"- **Quality Gap:** {round(gap * 100)}% (weaker side: {quality.get('weaker_side', 'N/A')})")

    lines.append("\n\n---")

    # Report metadata — word count and estimated reading time
    full_text = "\n".join(lines)
    word_count = len(full_text.split())
    reading_time = max(1, round(word_count / 200))  # ~200 wpm average
    lines.append(f"*Generated by Verdict AI Courtroom — {word_count:,} words, ~{reading_time} min read*")

    return "\n".join(lines)


def generate_json_report(session_data: dict) -> str:
    """Generate a JSON export of the full session data."""
    return json.dumps(session_data, indent=2, default=str)


# Domain-specific PDF color themes and subtitles
DOMAIN_PDF_THEMES = {
    "business": {"accent": (180, 142, 58), "subtitle": "Strategic Business Decision Analysis"},
    "financial": {"accent": (16, 185, 129), "subtitle": "Financial & Investment Analysis"},
    "legal": {"accent": (99, 102, 241), "subtitle": "Legal Risk & Compliance Assessment"},
    "medical": {"accent": (239, 68, 68), "subtitle": "Clinical Evidence Review"},
    "technology": {"accent": (59, 130, 246), "subtitle": "Technical Architecture Evaluation"},
    "hiring": {"accent": (168, 85, 247), "subtitle": "Talent Acquisition Decision Brief"},
    "strategic": {"accent": (245, 158, 11), "subtitle": "Strategic Planning Assessment"},
    "product": {"accent": (20, 184, 166), "subtitle": "Product Strategy Analysis"},
    "marketing": {"accent": (236, 72, 153), "subtitle": "Market Strategy Evaluation"},
}


def generate_pdf_report(session_data: dict) -> bytes:
    """Generate a domain-themed PDF report from session data using fpdf2.

    Each domain gets a unique color accent and subtitle in the PDF header,
    providing visually distinct reports for business, legal, medical, etc.
    """
    from fpdf import FPDF, XPos, YPos

    decision = session_data.get("decision", {})
    research = session_data.get("research_package", {})
    prosecutor = session_data.get("prosecutor_argument", {})
    defense = session_data.get("defense_argument", {})
    witnesses = session_data.get("witness_reports", [])
    verdict = session_data.get("verdict", {})
    synthesis = session_data.get("synthesis", {})
    output_format = session_data.get("output_format", "executive")
    domain = session_data.get("domain", "business")

    theme = DOMAIN_PDF_THEMES.get(domain, DOMAIN_PDF_THEMES["business"])
    accent = theme["accent"]

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    def safe_text(text: str, max_len: int = 500) -> str:
        """Sanitize text for PDF output."""
        if not text:
            return ""
        # Replace non-latin1 characters
        text = str(text)[:max_len]
        return text.encode("latin-1", errors="replace").decode("latin-1")

    # Title — domain accent color
    pdf.set_font("Helvetica", style="B", size=20)
    pdf.set_text_color(*accent)
    pdf.cell(0, 12, "VERDICT - AI Courtroom Analysis", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # Domain subtitle
    pdf.set_font("Helvetica", style="I", size=11)
    pdf.set_text_color(accent[0] // 2 + 60, accent[1] // 2 + 60, accent[2] // 2 + 60)
    pdf.cell(0, 7, theme["subtitle"], new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    # Metadata
    pdf.set_font("Helvetica", size=10)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 6, f"Decision: {safe_text(decision.get('question', 'N/A'), 120)}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.cell(0, 6, f"Date: {datetime.utcnow().strftime('%B %d, %Y')}  |  Format: {output_format.title()}  |  Domain: {domain.title()}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(4)

    # Divider — domain accent color
    pdf.set_draw_color(*accent)
    pdf.set_line_width(0.5)
    pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
    pdf.ln(6)

    def section_header(title: str):
        pdf.set_font("Helvetica", style="B", size=13)
        pdf.set_text_color(*accent)
        pdf.cell(0, 8, title, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        # Accent underline for section headers
        pdf.set_draw_color(*accent)
        pdf.set_line_width(0.3)
        pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 60, pdf.get_y())
        pdf.ln(3)

    def body_text(text: str, indent: int = 0):
        pdf.set_font("Helvetica", size=10)
        pdf.set_text_color(50, 50, 50)
        if indent:
            pdf.set_x(pdf.get_x() + indent)
        pdf.multi_cell(190 - indent, 5.5, safe_text(text, 600))
        pdf.ln(1)

    def bullet(text: str):
        pdf.set_font("Helvetica", size=10)
        pdf.set_text_color(80, 80, 80)
        pdf.set_x(pdf.get_x() + 6)
        pdf.multi_cell(183, 5, f"- {safe_text(text, 300)}")

    # 1. Verdict ruling (most important, put first for exec format)
    if verdict:
        ruling = verdict.get("ruling", "conditional").upper()
        ruling_color = (16, 185, 129) if ruling == "PROCEED" else (239, 68, 68) if ruling == "REJECT" else (245, 158, 11)
        section_header("THE RULING")
        pdf.set_font("Helvetica", style="B", size=18)
        pdf.set_text_color(*ruling_color)
        pdf.cell(0, 10, ruling, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(50, 50, 50)
        pdf.set_font("Helvetica", size=10)
        pdf.multi_cell(190, 5.5, safe_text(verdict.get("reasoning", ""), 800))
        pdf.ln(3)
        if verdict.get("key_factors"):
            pdf.set_font("Helvetica", style="B", size=10)
            pdf.set_text_color(80, 80, 80)
            pdf.cell(0, 6, "Key Factors:", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            for factor in verdict["key_factors"]:
                bullet(factor)
        pdf.ln(4)

    # 2. Research
    if research:
        section_header("RESEARCH ANALYSIS")
        if research.get("summary"):
            body_text(research["summary"])
        if research.get("key_data_points"):
            pdf.set_font("Helvetica", style="B", size=10)
            pdf.set_text_color(80, 80, 80)
            pdf.cell(0, 6, "Key Data Points:", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            for point in research["key_data_points"][:5]:
                bullet(str(point))
        pdf.ln(4)

    # 3. Prosecution
    if prosecutor:
        section_header("PROSECUTION (FOR)")
        if prosecutor.get("opening"):
            body_text(prosecutor["opening"])
        for i, claim in enumerate(prosecutor.get("claims", [])[:4], 1):
            pdf.set_font("Helvetica", style="B", size=10)
            pdf.set_text_color(239, 68, 68)
            pdf.cell(0, 6, f"Claim {i}: {safe_text(claim.get('statement', ''), 120)}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("Helvetica", size=9)
            pdf.set_text_color(100, 100, 100)
            pdf.multi_cell(190, 5, safe_text(claim.get("evidence", ""), 300))
            pdf.ln(1)
        pdf.ln(3)

    # 4. Defense
    if defense:
        section_header("DEFENSE (AGAINST)")
        if defense.get("opening"):
            body_text(defense["opening"])
        for i, claim in enumerate(defense.get("claims", [])[:4], 1):
            pdf.set_font("Helvetica", style="B", size=10)
            pdf.set_text_color(59, 130, 246)
            pdf.cell(0, 6, f"Claim {i}: {safe_text(claim.get('statement', ''), 120)}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("Helvetica", size=9)
            pdf.set_text_color(100, 100, 100)
            pdf.multi_cell(190, 5, safe_text(claim.get("evidence", ""), 300))
            pdf.ln(1)
        pdf.ln(3)

    # 5. Witnesses
    if witnesses:
        section_header("WITNESS REPORTS")
        for i, w in enumerate(witnesses, 1):
            w_verdict = w.get("verdict_on_claim", "inconclusive")
            vcolor = (16, 185, 129) if w_verdict == "sustained" else (239, 68, 68) if w_verdict == "overruled" else (245, 158, 11)
            pdf.set_font("Helvetica", style="B", size=10)
            pdf.set_text_color(*vcolor)
            pdf.cell(0, 6, f"Witness {i} ({w.get('witness_type', 'fact').title()}): {w_verdict.upper()}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font("Helvetica", size=9)
            pdf.set_text_color(80, 80, 80)
            pdf.multi_cell(190, 5, safe_text(w.get("resolution", ""), 400))
            pdf.ln(2)
        pdf.ln(3)

    # 6. Synthesis
    if synthesis:
        section_header("BATTLE-TESTED SYNTHESIS")
        if synthesis.get("improved_idea"):
            body_text(synthesis["improved_idea"])
        if synthesis.get("recommended_actions"):
            pdf.set_font("Helvetica", style="B", size=10)
            pdf.set_text_color(80, 80, 80)
            pdf.cell(0, 6, "Recommended Actions:", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            for i, action in enumerate(synthesis["recommended_actions"], 1):
                bullet(f"{i}. {action}")

    # Footer — domain accent divider
    pdf.ln(8)
    pdf.set_draw_color(*accent)
    pdf.set_line_width(0.5)
    pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
    pdf.ln(4)
    pdf.set_font("Helvetica", size=8)
    pdf.set_text_color(150, 150, 150)
    pdf.cell(0, 5, f"Generated by Verdict AI Courtroom  |  {theme['subtitle']}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    return bytes(pdf.output())


def generate_docx_report(session_data: dict) -> bytes:
    """Generate a formatted DOCX report from session data using python-docx."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    decision = session_data.get("decision", {})
    research = session_data.get("research_package", {})
    prosecutor = session_data.get("prosecutor_argument", {})
    defense = session_data.get("defense_argument", {})
    witnesses = session_data.get("witness_reports", [])
    verdict = session_data.get("verdict", {})
    synthesis = session_data.get("synthesis", {})
    output_format = session_data.get("output_format", "executive")
    domain = session_data.get("domain", "business")

    doc = Document()

    # Title
    title = doc.add_heading("VERDICT — AI Courtroom Analysis", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Decision: {decision.get('question', 'N/A')}\n").bold = True
    meta.add_run(f"Date: {datetime.utcnow().strftime('%B %d, %Y')}  |  ")
    meta.add_run(f"Format: {output_format.title()}  |  Domain: {domain.title()}")
    doc.add_paragraph("─" * 60)

    # Verdict
    if verdict:
        doc.add_heading("THE RULING", level=1)
        ruling = verdict.get("ruling", "conditional").upper()
        ruling_para = doc.add_paragraph()
        ruling_run = ruling_para.add_run(ruling)
        ruling_run.bold = True
        ruling_run.font.size = Pt(24)
        if ruling == "PROCEED":
            ruling_run.font.color.rgb = RGBColor(16, 185, 129)
        elif ruling == "REJECT":
            ruling_run.font.color.rgb = RGBColor(239, 68, 68)
        else:
            ruling_run.font.color.rgb = RGBColor(245, 158, 11)

        conf = round(verdict.get("confidence", 0.5) * 100)
        doc.add_paragraph(f"Confidence: {conf}%")
        if verdict.get("reasoning"):
            doc.add_paragraph(verdict["reasoning"])
        if verdict.get("key_factors"):
            doc.add_heading("Key Factors", level=2)
            for factor in verdict["key_factors"]:
                doc.add_paragraph(factor, style="List Bullet")

    # Research
    if research:
        doc.add_heading("RESEARCH ANALYSIS", level=1)
        if research.get("summary"):
            doc.add_paragraph(research["summary"])
        if research.get("key_data_points"):
            for point in research["key_data_points"][:5]:
                doc.add_paragraph(str(point), style="List Bullet")

    # Prosecution
    if prosecutor:
        doc.add_heading("PROSECUTION (FOR)", level=1)
        if prosecutor.get("opening"):
            doc.add_paragraph(prosecutor["opening"]).italic = True
        for i, claim in enumerate(prosecutor.get("claims", [])[:4], 1):
            p = doc.add_paragraph()
            p.add_run(f"Claim {i}: ").bold = True
            p.add_run(claim.get("statement", ""))
            doc.add_paragraph(f"Evidence: {claim.get('evidence', '')}")
            doc.add_paragraph(f"Confidence: {round(claim.get('confidence', 0.5) * 100)}%")

    # Defense
    if defense:
        doc.add_heading("DEFENSE (AGAINST)", level=1)
        if defense.get("opening"):
            doc.add_paragraph(defense["opening"]).italic = True
        for i, claim in enumerate(defense.get("claims", [])[:4], 1):
            p = doc.add_paragraph()
            p.add_run(f"Claim {i}: ").bold = True
            p.add_run(claim.get("statement", ""))
            doc.add_paragraph(f"Evidence: {claim.get('evidence', '')}")
            doc.add_paragraph(f"Confidence: {round(claim.get('confidence', 0.5) * 100)}%")

    # Witnesses
    if witnesses:
        doc.add_heading("WITNESS REPORTS", level=1)
        for i, w in enumerate(witnesses, 1):
            w_verdict = w.get("verdict_on_claim", "inconclusive").upper()
            doc.add_heading(f"Witness {i} ({w.get('witness_type', 'fact').title()}): {w_verdict}", level=2)
            doc.add_paragraph(w.get("resolution", ""))
            doc.add_paragraph(f"Confidence: {round(w.get('confidence', 0.5) * 100)}%")

    # Synthesis
    if synthesis:
        doc.add_heading("BATTLE-TESTED SYNTHESIS", level=1)
        if synthesis.get("improved_idea"):
            doc.add_paragraph(synthesis["improved_idea"])
        if synthesis.get("addressed_objections"):
            doc.add_heading("Objections Addressed", level=2)
            for obj in synthesis["addressed_objections"]:
                doc.add_paragraph(f"✓ {obj}", style="List Bullet")
        if synthesis.get("recommended_actions"):
            doc.add_heading("Recommended Actions", level=2)
            for i, action in enumerate(synthesis["recommended_actions"], 1):
                doc.add_paragraph(f"{i}. {action}")

    # Footer
    doc.add_paragraph("─" * 60)
    footer = doc.add_paragraph("Generated by Verdict AI Courtroom — Multi-Agent Adversarial Decision Analysis")
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
